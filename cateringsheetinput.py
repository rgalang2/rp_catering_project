from cateringsheet import info_list, magnum_rolls, banh_mi, flb
from docx import Document
from docx.shared import Pt
from tkinter import *
from tkinter import ttk

root = Tk()
root.title("Catering Sheet Maker")
mainframe = ttk.Frame(root, padding="40 40 40 40")
mainframe.grid(column=0, row=0, sticky=(N, W, E, S))
name = StringVar()
number = StringVar()
date = StringVar()
time = StringVar()
delivery = BooleanVar()
address = StringVar()
special_instructions = StringVar()
food_item = StringVar()

name_label = Label(mainframe, text = "Guest Name: ")
name_label.grid(row = 0, column = 0, sticky = W, pady = 2)
name_entry = Entry(mainframe, textvariable=name)
name_entry.grid(row = 0, column = 1, pady = 2)

number_label = Label(mainframe, text = "Phone Number: ")
number_label.grid(row = 1, column = 0, sticky = W, pady = 2)
number_entry = Entry(mainframe, textvariable=number)
number_entry.grid(row = 1, column = 1, pady = 2)

date_label = Label(mainframe, text = "Date: ")
date_label.grid(row = 2, column = 0, sticky = W, pady = 2)
date_entry = Entry(mainframe, textvariable=date)
date_entry.grid(row = 2, column = 1, pady = 2)

time_label = Label(mainframe, text = "Time Due: ")
time_label.grid(row = 3, column = 0, sticky = W, pady = 2)
time_entry = Entry(mainframe, textvariable=time)
time_entry.grid(row = 3, column = 1, pady = 2)

delivery_check = ttk.Checkbutton(mainframe, text = "Delivery", variable=delivery, onvalue = True, offvalue=False)
delivery_check.grid(row = 4, column = 0, sticky = W, pady = 2)

menu = ["Magnum Rolls", "Banh Mi", "Fully Loaded Bowl", "Pho", "Fried Rice", "Veggie Egg Rolls", "Shrimp and Pork Egg Rolls"]
item_combobox = ttk.Combobox(mainframe, textvariable = food_item, values = menu)
item_combobox.grid(row = 6, column = 0, sticky = W, pady = 2)
item_combobox.state(["readonly"])

item_frame = Frame(mainframe)
item_frame.grid(row=7, column=0, columnspan=2, sticky=(W, E))

size_var = StringVar()
quantity_var = StringVar()
protein_var = StringVar()
sauces_var = StringVar()
mayo_var = StringVar()

def clear_inputs():
	for widget in item_frame.winfo_children():
		widget.destroy()
		
def add_item_to_list():
	selection = food_item.get()
	qty = int(quantity_var.get())
	sz = size_var.get()
	protein = protein_var.get()
	
	if selection == "Magnum Rolls":
		sauces = [s.strip() for s in sauces_var.get().split(",") if s.strip()]
		magnum_rolls(sz, qty, protein, sauces)

	elif selection == "Banh Mi":
		mayo = mayo_var.get()
		banh_mi(sz, qty, protein, mayo)

	elif selection == "Fully Loaded Bowl":
		sauces = [s.strip() for s in sauces_var.get().split(",") if s.strip()]
		flb(sz, qty, protein, sauces)

	clear_inputs()

def item_selected(event):
	clear_inputs()
	Label(item_frame, text="Size:").grid(row=0, column=0, sticky=W)
	Entry(item_frame, textvariable=size_var).grid(row=0, column=1)
	Label(item_frame, text="Quantity:").grid(row=1, column=0, sticky=W)
	Entry(item_frame, textvariable=quantity_var).grid(row=1, column=1)
	Label(item_frame, text="Protein:").grid(row=2, column=0, sticky=W)
	Entry(item_frame, textvariable=protein_var).grid(row=2, column=1)

	if food_item.get() == "Magnum Rolls":
		Label(item_frame, text="Sauces (comma separated):").grid(row=3, column=0, sticky=W)
		Entry(item_frame, textvariable=sauces_var).grid(row=3, column=1)
	elif food_item.get() == "Banh Mi":
		Label(item_frame, text="Mayo Type:").grid(row=3, column=0, sticky=W)
		Entry(item_frame, textvariable=mayo_var).grid(row=3, column=1)
	elif food_item.get() == "Fully Loaded Bowl":
		Label(item_frame, text="Sauces (comma separated):").grid(row=3, column=0, sticky=W)
		Entry(item_frame, textvariable=sauces_var).grid(row=3, column=1)

	Button(item_frame, text="Add Item", command=add_item_to_list).grid(row=5, column=0, columnspan=2, pady=10)

item_combobox.bind("<<ComboboxSelected>>", item_selected)
"""
FOR KIET/SEB
-this part is user input. i have no idea how to make a gui so good luck ! 
 
-below, i have it so the user keeps giving items until they type in 'x'
i mainly used it for test runs, but there is probably a better way for them to input stuff. 

-i kinda started in the code above, but im a noob pls help

lmk if you have questions B)
"""

# number = input("Guest's phone number: ")
# date = input("Date due: ")
# time = input("Time Due: ")
# delivery = False
# address = ""
# special_instructions = ""
# if input("Is this a delivery?: ") == "yes":
# 	delivery = True
# 	address = input("What is the address?: ")
# if input("Any special instructions?: ") == "yes":
# 	special_instructions = input("Special instructions: ")


user_input = input("What item?: ")
while user_input != 'x':
	if user_input == "magnum rolls":
		quantity = int(input("Quantity?: "))
		size = input("What size?: ")
		protein = input("What protein?: ")
		sauces = []
		sauce = input('What sauce?: ')
		while sauce != "no":
			sauces.append(sauce)
			sauce = input('Any other sauces?: ')
	
		magnum_rolls(size, quantity, protein, sauces)
		user_input = input("What item?: ")

	if user_input == "banh mi":
		quantity = int(input("Quantity?: "))
		size = input("What size?: ")
		protein = input("What protein?: ")
		mayo = input("What mayo?: ")

		banh_mi(size, quantity, protein, mayo)
		user_input = input("What item?: ")
	
	if user_input == "flb":
		quantity = int(input("Quantity?: "))
		size = input("What size?: ")
		protein = input("What protein?: ")
		sauces = []
		sauce = input("What sauce?: ")

		while sauce != "no":
			sauces.append(sauce)
			sauce = input("Any other sauces?: ")
		
		flb(size, quantity, protein, sauces)
		user_input = input("What item?: ")
		

info_dict = {}
prep_sheet = {}
sauce_list = {}
"""
THIS PART IS REALLY MESSY BUT IT WORKS 
it loops through info_list and organizes it in a dict (info_dict)

info_dict = {keys are food items: values are a dict of {protein:quantity} }
"""
for info in info_list:
	"""
	#magnum roll indexes: [0: quantity,  1: protein, 2: # of containers, 
	3: sauces, 4: napkins, 5: box volume, 6: ramekins, 7: total people]
	"""
	if "magnum rolls" in info[0]:
		if "Magnum Rolls" not in info_dict.keys():
			info_dict["Magnum Rolls"] = {info[1]:int(info[0].strip(" magnum rolls"))}
		else:
			if info[1] in info_dict["Magnum Rolls"].keys():
				info_dict["Magnum Rolls"][info[1]] += int(info[0].strip(" magnum rolls"))
			else:
				info_dict["Magnum Rolls"][info[1]] = int(info[0].strip(" magnum rolls"))
		for s in info[3]:
			if s not in sauce_list.keys():
				sauce_list[s] = info[2]
			else:
				sauce_list[s] += info[2]
		if "Napkins" not in prep_sheet.keys():
			prep_sheet["Napkins"] = info[4]
		else:
			prep_sheet["Napkins"] += info[4]
		if "Catering Boxes" not in prep_sheet.keys():
			prep_sheet["Catering Boxes"] = info[5]
		else:
			prep_sheet["Catering Boxes"] += info[5]
		if "Ramekins" not in prep_sheet.keys():
			prep_sheet["Ramekins"] = info[6]
		else:
			prep_sheet["Ramekins"] += info[6]
		if "Plates" not in prep_sheet.keys():
			prep_sheet["Plates"] = info[7]*2
			while prep_sheet["Plates"] % 5 != 0:
				prep_sheet["Plates"] += 1
		else:
			prep_sheet["Plates"] += info[7]
			while prep_sheet["Plates"] % 5 != 0:
				prep_sheet["Plates"] += 1

	#banh mi indexes: [0: quantity, 1: protein, 2: mayo, 3: napkins, 4: total people, 5: box volume]
	elif "banh mi" in info[0]:
		if "Banh Mi" not in info_dict.keys():
			info_dict["Banh Mi"] = {f'{info[1]} w/ {info[2]}':int(info[0].strip(" banh mi"))}
		else:
			if f'{info[1]} w/ {info[2]}' in info_dict["Banh Mi"].keys():
				info_dict["Banh Mi"][f'{info[1]} w/ {info[2]}'] += int(info[0].strip(" banh mi"))
			else:
				info_dict["Banh Mi"][f'{info[1]} w/ {info[2]}'] = int(info[0].strip(" banh mi"))
		if "Napkins" not in prep_sheet.keys():
			prep_sheet["Napkins"] = info[3]
		else:
			prep_sheet["Napkins"] += info[3]
		if "Plates" not in prep_sheet.keys():
			prep_sheet["Plates"] = info[4]
		else:
			prep_sheet["Plates"] += info[4]
		if "Catering Boxes" not in prep_sheet.keys():
			prep_sheet["Catering Boxes"] = info[5]
		else:
			prep_sheet["Catering Boxes"] += info[5]

		"""
		flb list indexes: [0: # of tins, 1: size tin, 2: protein, 
		3: # of containers, 4:sauces, 5: napkins, 6: ramekins, 7: total people, 8: box volume]
		"""
	elif "tins" in info[0]:
		if "Fully Loaded Bowls" not in info_dict.keys():
			info_dict["Fully Loaded Bowls"] = {f'{info[2]} in {info[1]}':int(info[0].strip(" tins"))}
		else:
			if f'{info[2]} in {info[1]}' in info_dict["Fully Loaded Bowls"].keys():
				info_dict["Fully Loaded Bowls"][f'{info[2]} in {info[1]}'] += int(info[0].strip(" tins"))
			else:
				info_dict["Fully Loaded Bowls"][f'{info[2]} in {info[1]}'] = int(info[0].strip(" tins"))
		for s in info[4]:
			if s in sauce_list.keys():
				sauce_list[s] += info[3]
			else:
				sauce_list[s] = info[3]
		if "Napkins" not in prep_sheet.keys():
			prep_sheet["Napkins"] = info[5]
		else:
			prep_sheet["Napkins"] += info[5]
		if "Plates" not in prep_sheet.keys():
			prep_sheet["Plates"] = info[7]
		else:
			prep_sheet["Plates"] += info[7]
		if "Ramekins" not in prep_sheet.keys():
			prep_sheet["Ramekins"] = info[6]
		else:
			prep_sheet["Ramekins"] += info[6]
		if "large tins" in info[1]:
			if "Large Catering Tins + Lids" not in prep_sheet.keys():
				prep_sheet["Large Catering Tins + Lids"] = int(info[0].strip(" tins"))
			else:
				prep_sheet["Large Catering Tins + Lids"] += int(info[0].strip(" tins"))
		if "medium tins" in info[1]:
			if "Medium Catering Tins + Lids" not in prep_sheet.keys():
				prep_sheet["Medium Catering Tins + Lids"] = int(info[0].strip(" tins"))
			else:
				prep_sheet["Medium Catering Tins + Lids"] += int(info[0].strip(" tins"))
		if "Catering Boxes" not in prep_sheet.keys():
			prep_sheet["Catering Boxes"] = info[8]
		else:
			prep_sheet["Catering Boxes"] += info[8]
		if "Forks" not in prep_sheet.keys():
			prep_sheet["Forks"] = info[7]
		else:
			prep_sheet["Forks"] += info[7]
		if "Spoons" not in prep_sheet.keys():
			prep_sheet["Spoons"] = info[7]
		else:
			prep_sheet["Spoons"] += info[7]
	#fried rice 
	#pho
	#egg rolls
#making plates and utensils a multiple of 10
while prep_sheet["Plates"] % 10 != 0:
	prep_sheet["Plates"] += 1
if "Forks" in prep_sheet.keys(): 
	while prep_sheet["Forks"] % 10 != 0 and prep_sheet["Spoons"] % 10 != 0:
		prep_sheet["Forks"] +=1
		prep_sheet["Spoons"] +=1

prep_sheet["Catering Boxes"] = round(prep_sheet["Catering Boxes"], -1)//100
if prep_sheet["Catering Boxes"] == 0:
	prep_sheet["Catering Boxes"] = 1

print(info_list)
print(info_dict)

#constructs the word doc
document = Document()
style = document.styles['Normal']
font = style.font
font.name = "Arial"
font.size = Pt(13)
d = document.add_paragraph(f'Guest Name: {name} \nGuest\'s Phone #: \nDate: \nTime Due: \n')
d.add_run("Delivery: \nDelivered by: \nAddress: ")
document.add_paragraph("Special Instructions: ")
for items in info_dict:
	thing = document.add_paragraph(items, style="List Bullet")
	thing.paragraph_format.space_before = 1
	thing.paragraph_format.space_after = 1
	for a in info_dict.get(items):
		protein = document.add_paragraph(f'{info_dict[items].get(a)} {a}', style="List Bullet 2")
		protein.paragraph_format.space_after = 1
		protein.paragraph_format.space_before = 1
preparation = document.add_paragraph("Prep Sheet:")
preparation.paragraph_format.space_before = Pt(3)
preparation.paragraph_format.space_after = Pt(1)
for prep in prep_sheet:
	thing = document.add_paragraph(f'{prep}: {prep_sheet[prep]}', style= "List Bullet")
	thing.paragraph_format.space_after = Pt(1)
	thing.paragraph_format.space_before = Pt(1)
sauce_prep = document.add_paragraph("Sauces (24oz):")
sauce_prep.paragraph_format.space_before = Pt(3)
sauce_prep.paragraph_format.space_after = Pt(1)
for sasa in sauce_list:
	fuck = document.add_paragraph(f'{sasa}: {sauce_list[sasa]}', style="List Bullet")
	fuck.paragraph_format.space_before = Pt(1)
	fuck.paragraph_format.space_after = Pt(1)
document.save(f'{name}.docx')
print(f'{name}.docx saved!')

root.mainloop()
